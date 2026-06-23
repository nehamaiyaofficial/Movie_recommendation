from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("Intelligent_Movie_Recommendation_System_Report.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
BORDER = "D9E2F3"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(row.cells[idx])


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.autofit = False
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        header_cells[idx].text = header
        set_cell_shading(header_cells[idx], LIGHT_GRAY)
        for paragraph in header_cells[idx].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor(31, 77, 120)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
    set_table_widths(table, widths)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_code_block(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.right_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    for line in text.splitlines():
        run = paragraph.add_run(line + "\n")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(31, 77, 120)


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(12 if style_name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(6 if style_name != "Heading 3" else 4)


def add_footer(document: Document) -> None:
    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Intelligent Personalized Movie Recommendation System")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(89, 89, 89)


def build_document() -> None:
    document = Document()
    configure_styles(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(8)
    run = title.add_run("Intelligent Personalized Movie Recommendation System")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(18)
    run = subtitle.add_run("Using Content-Based Filtering, TF-IDF, Cosine Similarity, and User Preference Analysis")
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(89, 89, 89)

    add_table(
        document,
        ["Project Detail", "Description"],
        [
            ["Project Domain", "Machine Learning / Recommendation System"],
            ["Frontend", "Streamlit"],
            ["Programming Language", "Python"],
            ["Libraries", "Pandas, NumPy, Scikit-Learn, Requests, Streamlit"],
            ["Dataset Sources", "MovieLens public dataset plus curated regional Indian movie dataset"],
            ["Recommendation Type", "Content-Based Filtering with weighted user preference scoring"],
        ],
        [2.0, 4.5],
    )

    document.add_page_break()

    add_heading(document, "1. Introduction")
    document.add_paragraph(
        "Recommendation systems help users discover relevant items from large collections. "
        "This project recommends movies by analyzing content features and user inputs such as "
        "favorite movies, previously watched movies, current genre preference, preferred language, "
        "and preferred movie industry."
    )

    add_heading(document, "2. Problem Statement")
    document.add_paragraph(
        "Users often spend unnecessary time searching for a suitable movie across multiple languages "
        "and film industries. The objective of this project is to build an intelligent system that "
        "suggests personalized movie recommendations based on user preferences and movie metadata."
    )

    add_heading(document, "3. Objectives")
    add_bullets(
        document,
        [
            "Build a personalized content-based movie recommendation system.",
            "Use TF-IDF vectorization to convert movie metadata into numerical vectors.",
            "Use cosine similarity to identify movies similar to the user's preferences.",
            "Combine favorite movies, watched movies, genre, language, and industry preference.",
            "Provide explainable recommendations with similarity score and recommendation reason.",
            "Support Hollywood, Bollywood, Tollywood, Sandalwood, Kollywood, and Mollywood filtering.",
        ],
    )

    add_heading(document, "4. Literature Survey")
    document.add_paragraph(
        "Movie recommendation systems commonly use collaborative filtering, content-based filtering, "
        "or hybrid techniques. Collaborative filtering depends on user-user or user-item interaction "
        "data, while content-based filtering works with item metadata. Since this project focuses on "
        "movie metadata and direct user preference inputs, content-based filtering is the most suitable "
        "approach."
    )

    add_heading(document, "5. System Architecture")
    add_code_block(
        document,
        "User Input\n"
        "  -> Favorite Movies\n"
        "  -> Previously Watched Movies\n"
        "  -> Current Genre\n"
        "  -> Preferred Language and Industry\n"
        "  -> Feature Engineering\n"
        "  -> TF-IDF Vectorization\n"
        "  -> Cosine Similarity\n"
        "  -> Weighted Scoring\n"
        "  -> Top Movie Recommendations",
    )

    add_heading(document, "6. Dataset Description")
    document.add_paragraph(
        "The project uses a merged dataset. The MovieLens public dataset provides a broad collection "
        "of Hollywood/English movies with ratings and tags. A regional Indian movie dataset was added "
        "to support Hindi, Telugu, Kannada, Tamil, and Malayalam recommendations."
    )
    add_table(
        document,
        ["Column", "Purpose"],
        [
            ["title", "Movie name used for search and display"],
            ["genres", "Genre information used for mood and content matching"],
            ["overview", "Short description included in text features"],
            ["cast", "Actor information when available"],
            ["director", "Director information when available"],
            ["keywords", "Important terms used for similarity calculation"],
            ["vote_average", "Rating used for display and fallback ranking"],
            ["language", "Language filter such as Hindi, Telugu, Tamil, Kannada, Malayalam, or English"],
            ["industry", "Industry filter such as Bollywood, Tollywood, Kollywood, Sandalwood, Mollywood, or Hollywood"],
        ],
        [1.55, 4.95],
    )

    add_heading(document, "7. Methodology")
    add_heading(document, "7.1 Data Preprocessing", level=2)
    add_bullets(
        document,
        [
            "Missing text values are replaced with empty strings.",
            "Ratings are converted into numerical values.",
            "Movie titles are normalized to improve matching.",
            "Fuzzy matching is used for small spelling mistakes in user input.",
        ],
    )

    add_heading(document, "7.2 Feature Engineering", level=2)
    document.add_paragraph("A combined movie feature called tags is created by joining important metadata fields.")
    add_code_block(document, "tags = genres + cast + director + keywords + overview")

    add_heading(document, "7.3 TF-IDF Vectorization", level=2)
    document.add_paragraph(
        "TF-IDF converts movie text features into numerical vectors. It gives higher importance to "
        "terms that are meaningful for a movie and lower importance to very common words."
    )

    add_heading(document, "7.4 Cosine Similarity", level=2)
    document.add_paragraph(
        "Cosine similarity measures the similarity between two movie vectors. Movies with similar "
        "genre, keywords, cast, director, or overview receive higher similarity scores."
    )

    add_heading(document, "7.5 Weighted Recommendation Scoring", level=2)
    add_code_block(
        document,
        "Final Score = 0.4 * Favorite Movie Similarity\n"
        "            + 0.3 * Watch History Similarity\n"
        "            + 0.3 * Genre Match",
    )
    document.add_paragraph(
        "After calculating the score, the system filters recommendations by selected language and "
        "industry. If no typed movie is found in the dataset, the system still recommends movies "
        "using genre, language, industry, and rating."
    )

    add_heading(document, "8. Algorithms and Techniques Used")
    add_table(
        document,
        ["Technique", "Role in Project"],
        [
            ["Data Preprocessing", "Cleans and standardizes movie metadata"],
            ["Feature Engineering", "Creates the combined tags feature"],
            ["TF-IDF Vectorization", "Converts text into machine-readable vectors"],
            ["Cosine Similarity", "Finds movies similar to user-selected movies"],
            ["Content-Based Filtering", "Recommends movies based on movie features"],
            ["Weighted Scoring", "Combines long-term favorites, recent watches, and current genre"],
            ["Fuzzy Matching", "Handles minor spelling mistakes in movie names"],
            ["Language/Industry Filtering", "Restricts output to the user's preferred viewing language or industry"],
        ],
        [2.1, 4.4],
    )

    add_heading(document, "9. Implementation")
    add_numbered(
        document,
        [
            "User enters favorite movies, watched movies, current genre, language, and industry.",
            "The system searches and normalizes movie names.",
            "Movie metadata is combined into the tags feature.",
            "TF-IDF converts tags into vectors.",
            "Cosine similarity compares user-selected movies with all dataset movies.",
            "Weighted scoring ranks candidate movies.",
            "Language and industry filters are applied.",
            "The top recommendations are displayed in Streamlit with explanation and score.",
        ],
    )

    add_heading(document, "10. Results")
    document.add_paragraph(
        "The system displays recommended movies with name, genre, rating, language, industry, "
        "description, similarity score, and reason for recommendation."
    )
    add_table(
        document,
        ["Input Example", "Value"],
        [
            ["Favorite Movies", "Chhichhore, Dhurandhar, Race Gurram"],
            ["Previously Watched", "Maa Behen"],
            ["Current Genre", "Comedy"],
            ["Language", "Hindi"],
            ["Industry", "Bollywood"],
            ["Example Output", "Kuch Kuch Hota Hai, Munna Bhai M.B.B.S., Dil Chahta Hai, PK, 3 Idiots"],
        ],
        [2.0, 4.5],
    )

    add_heading(document, "11. Advantages")
    add_bullets(
        document,
        [
            "Works without user-user interaction data.",
            "Explains why a movie was recommended.",
            "Supports multiple Indian film industries and Hollywood.",
            "Handles spelling mistakes using fuzzy matching.",
            "Can be extended with larger datasets such as IMDb, TMDB, or Kaggle sources.",
        ],
    )

    add_heading(document, "12. Limitations")
    add_bullets(
        document,
        [
            "Recommendations depend on dataset quality and coverage.",
            "Collaborative behavior is not included because user interaction history is unavailable.",
            "Poster quality depends on optional TMDB API access.",
            "A complete all-language movie dataset may require authenticated APIs or licensed data sources.",
        ],
    )

    add_heading(document, "13. Future Scope")
    add_bullets(
        document,
        [
            "Add full TMDB API integration for live movie metadata.",
            "Add collaborative filtering if user ratings become available.",
            "Add user login and saved preference profiles.",
            "Add sentiment analysis using movie reviews.",
            "Deploy the Streamlit app online.",
        ],
    )

    add_heading(document, "14. Conclusion")
    document.add_paragraph(
        "The project successfully demonstrates an intelligent personalized movie recommendation "
        "system using content-based filtering and user preference analysis. The weighted scoring "
        "method combines long-term preference, recent behavior, current mood, language preference, "
        "and industry preference to generate useful and explainable recommendations."
    )

    add_heading(document, "15. Viva Questions")
    add_table(
        document,
        ["Question", "Answer"],
        [
            ["What type of recommendation system is this?", "It is a content-based recommendation system with weighted user preference analysis."],
            ["Why not collaborative filtering?", "Collaborative filtering requires user-user or user-item interaction data, which is not available."],
            ["What is TF-IDF?", "TF-IDF converts text into numerical vectors based on term importance."],
            ["What is cosine similarity?", "It measures similarity between two vectors based on the angle between them."],
            ["Why use weighted scoring?", "It combines favorite movies, recent watch history, and current genre preference."],
            ["Why add language and industry filters?", "They ensure the final output matches what the user wants to watch now."],
        ],
        [2.5, 4.0],
    )

    add_footer(document)
    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
