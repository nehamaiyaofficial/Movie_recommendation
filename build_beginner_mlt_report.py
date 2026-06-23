from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DATASET = Path("data/movies_clean.csv")
OUTPUT = Path("MLT_Movie_Recommendation_Project_Report.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
BORDER = "B4C6E7"


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def border_cell(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:color"), BORDER)


def add_table(document: Document, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        shade_cell(cell, LIGHT_GRAY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9.5)
                run.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    for values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
    for row in table.rows:
        for index, width in enumerate(widths):
            row.cells[index].width = Inches(width)
            border_cell(row.cells[index])
    document.add_paragraph().paragraph_format.space_after = Pt(0)


def add_heading(document: Document, text: str, level: int = 1) -> None:
    paragraph = document.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Number")
        paragraph.paragraph_format.space_after = Pt(4)
        paragraph.add_run(item)


def add_code(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.3)
    paragraph.paragraph_format.right_indent = Inches(0.3)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    for line in text.splitlines():
        run = paragraph.add_run(line + "\n")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor.from_string(DARK_BLUE)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_settings = {
        "Heading 1": (16, BLUE, 14, 7),
        "Heading 2": (13, BLUE, 10, 5),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_settings.items():
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_cover(document: Document) -> None:
    document.add_paragraph("\n\n")
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("INTELLIGENT PERSONALIZED\nMOVIE RECOMMENDATION SYSTEM")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string(DARK_BLUE)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(30)
    run = subtitle.add_run("Machine Learning Techniques Lab Project Report")
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(89, 89, 89)

    details = document.add_paragraph()
    details.alignment = WD_ALIGN_PARAGRAPH.CENTER
    details.paragraph_format.line_spacing = 1.5
    details.add_run("Submitted by\n").bold = True
    details.add_run("Name: ______________________________\n")
    details.add_run("USN / Roll Number: __________________\n")
    details.add_run("Class: 3rd Year B.Tech\n\n")
    details.add_run("Submitted to\n").bold = True
    details.add_run("Department of ________________________\n")
    details.add_run("College: _____________________________\n")
    details.add_run("Academic Year: 2025-2026")
    document.add_page_break()


def dataset_stats(data: pd.DataFrame) -> tuple[list[list[str]], list[list[str]]]:
    language_rows = [[language, str(count)] for language, count in data["language"].value_counts().items()]
    industry_rows = [[industry, str(count)] for industry, count in data["industry"].value_counts().items()]
    return language_rows, industry_rows


def build_report() -> None:
    data = pd.read_csv(DATASET)
    language_rows, industry_rows = dataset_stats(data)

    document = Document()
    configure_document(document)
    add_cover(document)

    add_heading(document, "Abstract")
    document.add_paragraph(
        "This project presents a beginner-level personalized movie recommendation system. The system "
        "asks the user for favorite movies, previously watched movies, current genre, preferred "
        "language, and movie industry. It uses movie information such as genre, keywords, cast, "
        "director, and overview to find similar movies. TF-IDF is used to convert text into numerical "
        "vectors, cosine similarity is used to compare movies, and a simple weighted formula is used "
        "to rank the final recommendations. The application is developed in Python with a Streamlit "
        "interface."
    )

    add_heading(document, "1. Introduction")
    document.add_paragraph(
        "A recommendation system helps a user select suitable items from a large collection. Popular "
        "applications use recommendation systems for movies, music, products, and videos. In this "
        "project, I developed a movie recommendation system as part of the Machine Learning Techniques "
        "laboratory. The main aim is to understand basic data preprocessing, feature extraction, vector "
        "similarity, and ranking methods."
    )

    add_heading(document, "2. Problem Statement")
    document.add_paragraph(
        "Users may find it difficult to choose a movie because many movies are available in different "
        "genres, languages, and film industries. The system should take simple preference inputs and "
        "recommend movies that are related to those preferences."
    )

    add_heading(document, "3. Objectives")
    add_bullets(
        document,
        [
            "To understand basic concepts of a content-based recommendation system.",
            "To clean and prepare movie data before applying an algorithm.",
            "To convert movie text features into numerical TF-IDF vectors.",
            "To compare movies using cosine similarity.",
            "To combine favorite movies, watch history, and current genre using weighted scoring.",
            "To filter recommendations by language and movie industry.",
            "To create a simple user interface using Streamlit.",
        ],
    )

    add_heading(document, "4. Tools and Technologies")
    add_table(
        document,
        ["Tool / Library", "Use in Project"],
        [
            ["Python", "Main programming language"],
            ["Pandas", "Reading, cleaning, merging, and processing CSV data"],
            ["NumPy", "Numerical array operations"],
            ["Scikit-Learn", "TF-IDF vectorization and cosine similarity"],
            ["Streamlit", "Frontend user interface"],
            ["Requests", "Optional TMDB poster/API requests"],
        ],
        [1.8, 4.7],
    )

    add_heading(document, "5. Dataset Used")
    document.add_paragraph(
        "The final dataset is stored in data/movies_clean.csv. It is prepared by combining the public "
        "MovieLens dataset with a regional Indian movie CSV. MovieLens mainly provides English and "
        "Hollywood movies, ratings, genres, and tags. The regional CSV adds Hindi, Telugu, Tamil, "
        "Kannada, and Malayalam movies."
    )
    document.add_paragraph(f"Total records in the processed dataset: {len(data):,} movies.")

    add_heading(document, "5.1 Dataset Columns", level=2)
    add_table(
        document,
        ["Column", "Meaning"],
        [
            ["title", "Name of the movie"],
            ["genres", "Movie genres such as Comedy, Action, Drama, or Thriller"],
            ["overview", "Short description of the movie"],
            ["cast", "Main actors, when available"],
            ["director", "Director name, when available"],
            ["keywords", "Important movie tags and related words"],
            ["vote_average", "Average movie rating on a 10-point scale"],
            ["poster_path", "Optional path used to display a poster"],
            ["language", "English, Hindi, Telugu, Tamil, Kannada, or Malayalam"],
            ["industry", "Hollywood, Bollywood, Tollywood, Kollywood, Sandalwood, or Mollywood"],
        ],
        [1.55, 4.95],
    )

    add_heading(document, "5.2 Language Distribution", level=2)
    add_table(document, ["Language", "Number of Movies"], language_rows, [3.2, 3.3])

    add_heading(document, "5.3 Industry Distribution", level=2)
    add_table(document, ["Industry", "Number of Movies"], industry_rows, [3.2, 3.3])

    add_heading(document, "6. Data Preprocessing")
    document.add_paragraph(
        "Data preprocessing is done before applying the recommendation algorithm. The following simple "
        "steps are performed in prepare_dataset.py and app.py."
    )
    add_numbered(
        document,
        [
            "Read movies.csv, ratings.csv, tags.csv, and the regional movie CSV using Pandas.",
            "Merge rating and tag information using the movieId column.",
            "Replace missing text values with empty strings.",
            "Replace missing ratings with the mean rating.",
            "Convert MovieLens ratings from a 0-5 scale to a 0-10 scale.",
            "Replace the genre separator | with a space.",
            "Add language and industry columns.",
            "Remove duplicate movie records.",
            "Normalize movie titles by converting them to lowercase and removing release years.",
        ],
    )

    add_heading(document, "6.1 Example of Preprocessing", level=2)
    add_table(
        document,
        ["Before Processing", "After Processing"],
        [
            ["Action|Comedy|Drama", "Action Comedy Drama"],
            ["The Matrix (1999)", "the matrix"],
            ["MovieLens rating = 4.2", "Displayed rating = 8.4"],
            ["Missing cast or director", "Empty string"],
        ],
        [3.25, 3.25],
    )

    add_heading(document, "7. Data Processing and Feature Engineering")
    document.add_paragraph(
        "The model cannot directly understand movie descriptions as text. Therefore, useful text fields "
        "are combined into one feature called tags."
    )
    add_code(document, "tags = genres + cast + director + keywords + overview")
    document.add_paragraph(
        "For example, a movie may be represented by words such as action, comedy, friendship, college, "
        "Allu Arjun, and Telugu. This combined text gives a simple content profile for each movie."
    )

    add_heading(document, "8. Machine Learning Techniques Used")
    add_heading(document, "8.1 TF-IDF Vectorization", level=2)
    document.add_paragraph(
        "TF-IDF means Term Frequency-Inverse Document Frequency. It converts the tags text into numerical "
        "vectors. Important words receive higher values, while very common words receive lower values. "
        "TF-IDF is an NLP-based feature extraction method, but it is used here as part of the machine "
        "learning recommendation pipeline."
    )
    add_code(
        document,
        "vectorizer = TfidfVectorizer(stop_words='english', max_features=7000)\n"
        "vectors = vectorizer.fit_transform(data['tags'])",
    )

    add_heading(document, "8.2 Cosine Similarity", level=2)
    document.add_paragraph(
        "Cosine similarity compares two movie vectors. A higher value means the movies have more similar "
        "content. A value near 1 represents high similarity, and a value near 0 represents low similarity."
    )
    add_code(document, "similarity = cosine_similarity(vectors)")

    add_heading(document, "8.3 Content-Based Filtering", level=2)
    document.add_paragraph(
        "Content-based filtering recommends a movie by comparing its features with movies already liked "
        "or watched by the user. It does not require ratings from many different users."
    )

    add_heading(document, "8.4 Fuzzy Title Matching", level=2)
    document.add_paragraph(
        "Fuzzy matching is used to handle small spelling mistakes. For example, race guram can match "
        "Race Gurram, and dhurandar can match Dhurandhar. This is a supporting data-matching technique "
        "and not the main recommendation algorithm."
    )

    add_heading(document, "9. Recommendation Algorithm")
    document.add_paragraph("The recommendation process is performed using the following steps:")
    add_numbered(
        document,
        [
            "Accept favorite movies, watched movies, genre, language, and industry from the user.",
            "Find matching movie titles in the dataset.",
            "Calculate average similarity with favorite movies.",
            "Calculate average similarity with previously watched movies.",
            "Check whether each candidate movie matches the selected genre.",
            "Calculate the weighted final score.",
            "Remove movies already entered by the user.",
            "Filter candidates by language and industry.",
            "Sort movies by final score and display the top recommendations.",
        ],
    )

    add_heading(document, "9.1 Weighted Scoring Formula", level=2)
    add_code(
        document,
        "Final Score = 0.4 * Favorite Movie Similarity\n"
        "            + 0.3 * Watch History Similarity\n"
        "            + 0.3 * Genre Match",
    )
    add_table(
        document,
        ["Input", "Weight", "Reason"],
        [
            ["Favorite movies", "40%", "Represents long-term interest"],
            ["Previously watched movies", "30%", "Represents recent viewing interest"],
            ["Current genre", "30%", "Represents the user's current mood"],
        ],
        [2.3, 1.0, 3.2],
    )

    add_heading(document, "9.2 Simple Pseudocode", level=2)
    add_code(
        document,
        "INPUT favorite_movies, watched_movies, genre, language, industry\n"
        "LOAD processed movie dataset\n"
        "CREATE TF-IDF vectors from movie tags\n"
        "CALCULATE cosine similarity\n"
        "FOR every candidate movie:\n"
        "    calculate favorite similarity\n"
        "    calculate watched similarity\n"
        "    calculate genre match\n"
        "    calculate weighted final score\n"
        "FILTER by language and industry\n"
        "SORT by final score\n"
        "DISPLAY top recommendations",
    )

    add_heading(document, "10. System Workflow")
    add_code(
        document,
        "User Input\n"
        "   -> Title Matching and Data Cleaning\n"
        "   -> Feature Engineering\n"
        "   -> TF-IDF Vectorization\n"
        "   -> Cosine Similarity\n"
        "   -> Weighted Scoring\n"
        "   -> Language and Industry Filtering\n"
        "   -> Top Movie Recommendations",
    )

    add_heading(document, "11. Sample Input and Output")
    add_table(
        document,
        ["User Input", "Example"],
        [
            ["Favorite movies", "Chhichhore, Dhurandhar, Race Gurram"],
            ["Previously watched", "Maa Behen"],
            ["Current genre", "Comedy"],
            ["Language", "Hindi"],
            ["Industry", "Bollywood"],
        ],
        [2.0, 4.5],
    )
    document.add_paragraph("Possible recommendations from the current dataset include:")
    add_bullets(
        document,
        [
            "Kuch Kuch Hota Hai",
            "Munna Bhai M.B.B.S.",
            "Dil Chahta Hai",
            "PK",
            "3 Idiots",
        ],
    )

    add_heading(document, "12. Why Train-Test Split Is Not Used")
    document.add_paragraph(
        "This project does not use a supervised prediction model with labeled output classes. It is an "
        "unsupervised content-based recommendation approach that calculates similarity between movie "
        "features. Therefore, a normal train-test split and accuracy score are not required for the "
        "current implementation. The results are checked using sample user inputs and by observing "
        "whether the recommended movies match the selected preferences."
    )

    add_heading(document, "13. Results")
    document.add_paragraph(
        "The system successfully accepts user preferences and displays ranked movie recommendations. "
        "The recommendations change when the user changes the favorite movies, watch history, genre, "
        "language, or industry. The system also gives a short explanation and similarity score for each "
        "recommended movie."
    )

    add_heading(document, "14. Advantages")
    add_bullets(
        document,
        [
            "The method is simple to understand and suitable for a beginner MLT project.",
            "It does not require personal information or ratings from many users.",
            "It supports multiple languages and movie industries.",
            "The weighted score is easy to explain during viva.",
            "The Streamlit interface makes the model easy to demonstrate.",
        ],
    )

    add_heading(document, "15. Limitations")
    add_bullets(
        document,
        [
            "Recommendation quality depends on the movies available in the dataset.",
            "The regional dataset contains fewer movies than the English MovieLens dataset.",
            "The model may recommend similar movies repeatedly because it does not learn from long-term feedback.",
            "Poster display requires an optional TMDB API key.",
            "The system does not currently use collaborative filtering.",
        ],
    )

    add_heading(document, "16. Future Scope")
    add_bullets(
        document,
        [
            "Use a larger authenticated TMDB, IMDb, or Kaggle dataset.",
            "Add collaborative filtering using real user ratings.",
            "Add login and save each user's recommendation history.",
            "Collect user feedback such as Like and Dislike.",
            "Deploy the Streamlit application online.",
        ],
    )

    add_heading(document, "17. Conclusion")
    document.add_paragraph(
        "This project helped me understand the basic steps involved in building a recommendation system. "
        "I learned how to clean a dataset, combine useful features, convert text into TF-IDF vectors, "
        "compare vectors using cosine similarity, and rank results using a weighted score. The final "
        "application is simple, explainable, and suitable for demonstrating basic Machine Learning "
        "Techniques concepts."
    )

    add_heading(document, "18. Viva Questions and Short Answers")
    add_table(
        document,
        ["Question", "Short Answer"],
        [
            ["What type of system is this?", "A content-based movie recommendation system."],
            ["What is TF-IDF?", "It converts text into numerical vectors based on word importance."],
            ["What is cosine similarity?", "It measures similarity between two numerical vectors."],
            ["Why is collaborative filtering not used?", "The project does not have sufficient user-user rating data."],
            ["Why use weighted scoring?", "It combines favorite movies, watched movies, and current genre."],
            ["Is TF-IDF an NLP topic?", "Yes. It is used here as text feature extraction inside the ML recommendation pipeline."],
            ["Why is no train-test split used?", "The model ranks items by similarity and is not a supervised classifier."],
        ],
        [2.4, 4.1],
    )

    add_heading(document, "19. References")
    add_numbered(
        document,
        [
            "GroupLens MovieLens Dataset: https://grouplens.org/datasets/movielens/",
            "Scikit-Learn TF-IDF Documentation: https://scikit-learn.org/stable/modules/generated/sklearn.feature_extraction.text.TfidfVectorizer.html",
            "Scikit-Learn Cosine Similarity Documentation: https://scikit-learn.org/stable/modules/generated/sklearn.metrics.pairwise.cosine_similarity.html",
            "Streamlit Documentation: https://docs.streamlit.io/",
        ],
    )

    for section in document.sections:
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("MLT Lab Project - Intelligent Movie Recommendation System")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(89, 89, 89)

    document.save(OUTPUT)
    print(f"Created {OUTPUT} with {len(data)} dataset records documented.")


if __name__ == "__main__":
    build_report()
